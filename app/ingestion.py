"""
Document ingestion module for parsing and chunking PDFs.
"""
import fitz  # PyMuPDF
import nltk
import tiktoken
from dataclasses import dataclass
import uuid
from pathlib import Path
from typing import List, Dict, Any
import docx
import requests
from bs4 import BeautifulSoup

@dataclass
class Chunk:
    """Represents a chunk of text from a document."""
    chunk_id: str          # uuid4 string
    doc_name: str          # filename without extension
    page_number: int       # 1-indexed, page of first sentence in chunk
    text: str              # raw chunk text
    token_count: int       # tiktoken count

def extract_pages(pdf_path: Path) -> List[Dict[str, Any]]:
    """
    Extracts text page by page from a given PDF file.
    
    Args:
        pdf_path: Path to the PDF file.
        
    Returns:
        List of dictionaries containing page_number, text, and doc_name.
    """
    if not pdf_path.exists():
        raise FileNotFoundError(f"File not found: {pdf_path}")
        
    doc_name = pdf_path.stem
    doc = fitz.open(str(pdf_path))
    pages = []
    
    # Extract tables using Camelot
    tables_by_page = {}
    try:
        import camelot
        # Suppress warnings and try to read tables
        tables = camelot.read_pdf(str(pdf_path), pages='all', flavor='stream', suppress_stdout=True)
        for table in tables:
            page_num = table.page
            # Convert to markdown
            md_table = table.df.to_markdown(index=False, header=False)
            if page_num not in tables_by_page:
                tables_by_page[page_num] = []
            tables_by_page[page_num].append(md_table)
    except Exception as e:
        print(f"Camelot table extraction skipped/failed (Ghostscript may be missing): {e}")
    
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        text = page.get_text("text")
        
        camelot_page_num = page_num + 1
        
        # Append any tables found on this page with a [TABLE] identifier
        if camelot_page_num in tables_by_page:
            for md_table in tables_by_page[camelot_page_num]:
                text += f"\\n\\n[TABLE]\\n{md_table}\\n"
        
        if len(text.strip()) < 20:
            print(f"Warning: Skipping page {camelot_page_num} of {doc_name} (insufficient text).")
            continue
            
        pages.append({
            "page_number": camelot_page_num,
            "text": text,
            "doc_name": doc_name
        })
        
    doc.close()
    return pages

def semantic_chunk(pages: List[Dict[str, Any]], max_tokens: int = 400, overlap_tokens: int = 80) -> List[Chunk]:
    """
    Chunks extracted pages into semantic chunks based on sentence boundaries.
    
    Args:
        pages: List of extracted page dictionaries.
        max_tokens: Maximum tokens per chunk.
        overlap_tokens: Number of overlapping tokens between consecutive chunks.
        
    Returns:
        List of Chunk objects.
    """
    nltk.download("punkt", quiet=True)
    nltk.download("punkt_tab", quiet=True)
    
    try:
        encoder = tiktoken.get_encoding("cl100k_base")
    except Exception as e:
        raise RuntimeError(f"Failed to load tiktoken encoder: {e}")
        
    all_sentences = []
    for page in pages:
        doc_name = page["doc_name"]
        page_num = page["page_number"]
        text = page["text"]
        
        sentences = nltk.sent_tokenize(text)
        for s in sentences:
            if not s.strip():
                continue
            token_count = len(encoder.encode(s))
            all_sentences.append({
                "text": s,
                "token_count": token_count,
                "doc_name": doc_name,
                "page_num": page_num
            })
            
    chunks = []
    idx = 0
    while idx < len(all_sentences):
        chunk_start_idx = idx
        current_chunk_sentences = []
        current_token_count = 0
        first_sentence = all_sentences[idx]
        
        while idx < len(all_sentences) and current_token_count + all_sentences[idx]["token_count"] <= max_tokens:
            current_chunk_sentences.append(all_sentences[idx]["text"])
            current_token_count += all_sentences[idx]["token_count"]
            idx += 1
            
        if not current_chunk_sentences and idx < len(all_sentences):
            # Single sentence exceeds max_tokens, add it to avoid infinite loop
            current_chunk_sentences.append(all_sentences[idx]["text"])
            current_token_count += all_sentences[idx]["token_count"]
            idx += 1
            
        chunk_text = " ".join(current_chunk_sentences)
        chunks.append(Chunk(
            chunk_id=str(uuid.uuid4()),
            doc_name=first_sentence["doc_name"],
            page_number=first_sentence["page_num"],
            text=chunk_text,
            token_count=current_token_count
        ))
        
        if idx >= len(all_sentences):
            break
            
        # Backtrack for overlap
        overlap_count = 0
        backtrack_idx = idx - 1
        while backtrack_idx > chunk_start_idx:
            overlap_count += all_sentences[backtrack_idx]["token_count"]
            if overlap_count >= overlap_tokens:
                break
            backtrack_idx -= 1
            
        # Ensure we always make progress and overlap is properly captured
        idx = max(chunk_start_idx + 1, backtrack_idx)
        
    return chunks

def extract_docx(docx_path: Path) -> List[Dict[str, Any]]:
    if not docx_path.exists():
        raise FileNotFoundError(f"File not found: {docx_path}")
        
    doc_name = docx_path.stem
    doc = docx.Document(docx_path)
    pages = []
    
    text = "\\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_data:
                text += "\\n" + " | ".join(row_data)
                
    if text.strip():
        pages.append({
            "page_number": 1,
            "text": text,
            "doc_name": doc_name
        })
    return pages

def extract_text(txt_path: Path) -> List[Dict[str, Any]]:
    if not txt_path.exists():
        raise FileNotFoundError(f"File not found: {txt_path}")
        
    doc_name = txt_path.stem
    text = txt_path.read_text(encoding='utf-8', errors='ignore')
    
    pages = []
    if text.strip():
        pages.append({
            "page_number": 1,
            "text": text,
            "doc_name": doc_name
        })
    return pages

def extract_url(url: str) -> List[Dict[str, Any]]:
    try:
        response = requests.get(url, timeout=10)
        response.raise_for_status()
    except Exception as e:
        print(f"Error fetching URL {url}: {e}")
        return []
        
    soup = BeautifulSoup(response.text, 'html.parser')
    
    for element in soup(["nav", "footer", "header", "script", "style", "aside"]):
        element.extract()
        
    text = soup.get_text(separator='\\n', strip=True)
    
    from urllib.parse import urlparse
    parsed = urlparse(url)
    doc_name = f"URL_{parsed.netloc}{parsed.path}".replace("/", "_")
    if doc_name.endswith("_"):
        doc_name = doc_name[:-1]
        
    pages = []
    if text.strip():
        pages.append({
            "page_number": 1,
            "text": text,
            "doc_name": doc_name
        })
    return pages

def ingest_file(file_path: Path) -> List[Chunk]:
    """
    Ingests a single file (PDF, DOCX, TXT) and returns semantic chunks.
    """
    suffix = file_path.suffix.lower()
    if suffix == '.pdf':
        pages = extract_pages(file_path)
    elif suffix == '.docx':
        pages = extract_docx(file_path)
    elif suffix in ['.txt', '.md', '.csv']:
        pages = extract_text(file_path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")
        
    chunks = semantic_chunk(pages)
    print(f"Ingested {file_path.name}: {len(pages)} 'pages', {len(chunks)} chunks")
    return chunks

def ingest_url(url: str) -> List[Chunk]:
    """
    Ingests a website URL and returns semantic chunks.
    """
    pages = extract_url(url)
    if not pages:
        return []
    chunks = semantic_chunk(pages)
    print(f"Ingested URL {url}: {len(chunks)} chunks")
    return chunks

def ingest_multiple(file_paths: List[Path]) -> List[Chunk]:
    """
    Ingests multiple files and returns an aggregated list of chunks.
    """
    all_chunks = []
    for file_path in file_paths:
        chunks = ingest_file(file_path)
        all_chunks.extend(chunks)
    print(f"Total chunks ingested: {len(all_chunks)}")
    return all_chunks
